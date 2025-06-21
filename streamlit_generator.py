import os
import streamlit as st
import json
from google.oauth2.service_account import Credentials

st.set_page_config(page_title="CRM генератор", layout="centered")

# === НАСТРОЙКИ ПАПОК ===
TEMPLATE_DIR = "templates"
OUTPUT_DIR = "output"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# === ИНИЦИАЛИЗАЦИЯ СЕССИИ ===
if 'page' not in st.session_state:
    st.session_state['page'] = 'main_menu'
if 'form_data' not in st.session_state:
    st.session_state['form_data'] = {}

if st.session_state['page'] == 'main_menu':
    st.title("👋 Добро пожаловать в Личный кабинет менеджера Leads-Solver")

    st.subheader("Выберите действие:")
    col1, col2 = st.columns(2)

    with col1:
        if st.button("📄 Подготовить документы клиенту"):
            st.session_state['page'] = 'select_parameters'

    with col2:
        if st.button("📋 Реестр неоплаченных счетов"):
            st.session_state['page'] = 'unpaid_registry'

elif st.session_state['page'] == 'select_parameters':
    st.title("📄 Подготовка документов клиенту")

    with st.form("doc_parameters_form"):
        our_company = st.selectbox("📌 От какого юрлица готовится документ?", ["ООО Клиентология", "ИП Матвейчук С.Р."])
        payer_type = st.selectbox("👤 Кто плательщик?", ["ООО", "ИП", "Физлицо"])
        service_type = st.selectbox("💼 Тип услуги:", [
            "Оплата за номера",
            "Номера с КЦ без гарантии",
            "Номера с КЦ с гарантией",
            "Оплата за Лид"
        ])
        doc_type = st.selectbox("📄 Какие документы?", [
            "Счёт и договор",
            "Только счёт"
        ])

        submitted = st.form_submit_button("Продолжить")
        if submitted:
            st.session_state['combo'] = (our_company, payer_type, service_type, doc_type)
            st.session_state['page'] = 'fill_fields_placeholder'

    if st.button("🔙 Назад в меню"):
        st.session_state['page'] = 'main_menu'

elif st.session_state['page'] == 'fill_fields_placeholder':
    import re
    from docx import Document
    from datetime import datetime
    from docxtpl import DocxTemplate
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
    from num2words import num2words
    import gspread
    import json
    from google.oauth2.service_account import Credentials

    def extract_ordered_variables_from_docx(doc_path):
        doc = Document(doc_path)
        seen = set()
        ordered_vars = []

        def extract_from_text(text):
            for match in re.findall(r"{{(.*?)}}", text):
                if match not in seen:
                    seen.add(match)
                    ordered_vars.append(match)

        for p in doc.paragraphs:
            extract_from_text(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    extract_from_text(cell.text)

        return ordered_vars

    with open("template_map.json", encoding="utf-8") as f:
        template_map = json.load(f)
    with open("field_labels.json", encoding="utf-8") as f:
        field_labels = json.load(f)
    with open("responsible_list.json", encoding="utf-8") as f:
        responsible_data = json.load(f)
    responsible_names = [item["name"] for item in responsible_data]

    selected_combo = st.session_state['combo']
    selected_entry = next((item for item in template_map if (
        item['our_company'] == selected_combo[0] and
        item['payer_type'] == selected_combo[1] and
        item['service_type'] == selected_combo[2] and
        item['doc_type'] == selected_combo[3]
    )), None)

    if not selected_entry:
        st.error("❌ Не найдена запись в template_map.json под выбранную комбинацию.")
    else:
        # Только рендерим форму, если мы на этой странице
        if st.session_state['page'] == 'fill_fields_placeholder':
            st.title("📝 Заполнение данных по шаблону")
            ordered_vars = []
            seen_vars = set()

            for template_file in selected_entry['template_paths']:
                path = os.path.join(TEMPLATE_DIR, template_file)
                if not os.path.exists(path):
                    st.warning(f"⚠️ Файл не найден: {template_file}")
                else:
                    vars_from_template = extract_ordered_variables_from_docx(path)
                    for var in vars_from_template:
                        if var not in seen_vars and not var.endswith("_words"):
                            seen_vars.add(var)
                            ordered_vars.append(var)

            input_values = {}

            with st.form("fill_form"):
                st.write(f"📄 Шаблоны: {', '.join(selected_entry['template_paths'])}")
                for var in ordered_vars:
                    label = field_labels.get(var, f"{{{{{var}}}}}")
                    value = st.text_input(label, value=st.session_state['form_data'].get(var, ""))
                    input_values[var] = value

                st.subheader("📎 Дополнительно для учёта")
                input_values["deal_link"] = st.text_input("🔗 Ссылка на сделку из Битрикс", value=st.session_state['form_data'].get("deal_link", ""))
                input_values["deal_type"] = st.selectbox("📌 Тип сделки", ["Новый", "Пролонгация"])
                input_values["responsible"] = st.selectbox("👤 Ответственный", responsible_names)

                submitted = st.form_submit_button("➡️ Скачать готовые документы")
                if submitted:
                    try:
                        scopes = ["https://www.googleapis.com/auth/spreadsheets"]

                        secret_json = st.secrets["gcp_service_account"]["json"]
                        service_account_info = json.loads(secret_json)
                        service_account_info["private_key"] = service_account_info["private_key"].replace('\\n', '\n')
                        credentials = Credentials.from_service_account_info(service_account_info, scopes=scopes)

                        client = gspread.authorize(credentials)
                        sh = client.open_by_url("https://docs.google.com/spreadsheets/d/1AeW7yFTp2KIVPoDoGgouvLRNkf80pLIyz-I9gIeQKL4/edit/")
                        worksheet = sh.worksheet("Реестр не оплаченных счетов")

                        row = [
                            input_values.get("deal_link", ""),
                            input_values.get("deal_type", ""),
                            input_values.get("responsible", ""),
                            input_values.get("total_amount", ""),
                            input_values.get("invoice_date", ""),
                            input_values.get("contract_number", ""),
                            input_values.get("payer_fio", "") or input_values.get("client_short_name", "")
                        ]
                        worksheet.append_row(row)

                        def upload_to_gdrive(filepath, filename):
                            drive_scopes = ["https://www.googleapis.com/auth/drive"]
                            drive_service_account_info = service_account_info.copy()
                            drive_credentials = Credentials.from_service_account_info(drive_service_account_info, scopes=drive_scopes)
                            drive_service = build("drive", "v3", credentials=drive_credentials)
                            file_metadata = {"name": filename, "parents": ["1z-b3pc71PMxjeU9tgwmIgjIKYLUYaEPM"]}
                            media = MediaFileUpload(filepath, resumable=True)
                            drive_service.files().create(body=file_metadata, media_body=media, fields="id").execute()

                        st.session_state['generated_files'] = []

                        for template_file in selected_entry['template_paths']:
                            tpl_path = os.path.join(TEMPLATE_DIR, template_file)
                            doc = DocxTemplate(tpl_path)
                            context = input_values.copy()

                            for key in list(context.keys()):
                                if key.endswith("_numeric"):
                                    try:
                                        number = float(context[key])
                                        if number.is_integer():
                                            number = int(number)
                                        context[key.replace("_numeric", "_words")] = num2words(number, lang='ru')
                                    except:
                                        context[key.replace("_numeric", "_words")] = "[ошибка]"

                            doc.render(context)

                            prefix = "СЧЁТ" if "счет" in template_file.lower() or "счёт" in template_file.lower() else "ДОГОВОР"
                            postfix = context.get("contract_number", "без_номера")
                            filename = f"{prefix}-{postfix}.docx"
                            full_path = os.path.join(OUTPUT_DIR, filename)
                            doc.save(full_path)
                            upload_to_gdrive(full_path, filename)
                            st.session_state['generated_files'].append((filename, full_path))

                        st.session_state['form_data'] = input_values
                        st.session_state['page'] = 'document_download'

                    except Exception as e:
                        st.error(f"⚠️ Ошибка при записи в Google Sheets: {e}")

        if st.button("🔙 Назад"):
            st.session_state['page'] = 'select_parameters'

elif st.session_state['page'] == 'unpaid_registry':
    st.title("📋 Реестр неоплаченных счетов")

    try:
        import gspread
        import pandas as pd

        scopes = ["https://www.googleapis.com/auth/spreadsheets.readonly"]
        secret_json = st.secrets["gcp_service_account"]["json"]
        service_account_info = json.loads(secret_json)
        service_account_info["private_key"] = service_account_info["private_key"].replace('\\n', '\n')
        credentials = Credentials.from_service_account_info(service_account_info, scopes=scopes)
        client = gspread.authorize(credentials)

        sh = client.open_by_url("https://docs.google.com/spreadsheets/d/1AeW7yFTp2KIVPoDoGgouvLRNkf80pLIyz-I9gIeQKL4/edit/")
        worksheet = sh.worksheet("Реестр не оплаченных счетов")
        values = worksheet.get_all_values()

        if not values or len(values) < 2:
            st.error("Лист пустой или содержит недостаточно данных для обработки")
            df = None
        else:
            df = pd.DataFrame(values[1:], columns=values[0])

        if df is not None and not df.empty:
            st.markdown("### 🔍 Фильтры")

            col1, col2 = st.columns(2)
            with col1:
                selected_resp = st.selectbox("👤 Ответственный", ["Все"] + sorted(df["Ответственный"].dropna().unique()))
            with col2:
                selected_type = st.selectbox("📄 Тип договора", ["Все"] + sorted(df["Тип договора"].dropna().unique()))

            filtered_df = df.copy()
            if selected_resp != "Все":
                filtered_df = filtered_df[filtered_df["Ответственный"] == selected_resp]
            if selected_type != "Все":
                filtered_df = filtered_df[filtered_df["Тип договора"] == selected_type]

            if filtered_df.empty:
                st.info("Нет подходящих записей.")
            else:
                styled_html = """
<style>
    table { width: 100%; border-collapse: collapse; margin-top: 1rem; }
    th, td { padding: 12px 16px; text-align: left; border-bottom: 1px solid #ddd; font-size: 15px; }
    th { background-color: #f3f3f3; font-weight: bold; }
    tr:hover { background-color: #f9f9f9; }
    a { color: #0366d6; text-decoration: none; }
    a:hover { text-decoration: underline; }
</style>
"""
                table_html = "<table><thead><tr>"
                for col in filtered_df.columns:
                    table_html += f"<th>{col}</th>"
                table_html += "</tr></thead><tbody>"

                for _, row in filtered_df.iterrows():
                    table_html += "<tr>"
                    for col in filtered_df.columns:
                        val = row[col]
                        if isinstance(val, str) and val.startswith("http"):
                            val = f'<a href="{val}" target="_blank">Открыть сделку</a>'
                        table_html += f"<td>{val}</td>"
                    table_html += "</tr>"

                table_html += "</tbody></table>"
                st.markdown(styled_html + table_html, unsafe_allow_html=True)

    except Exception as e:
        st.error(f"❌ Не удалось загрузить таблицу: {e}")

    if st.button("🔙 Назад в меню"):
        st.session_state['page'] = 'main_menu'

elif st.session_state['page'] == 'document_download':
    st.title("✅ Документы успешно созданы")
    st.success("Документы загружены на Google Диск и готовы к скачиванию:")

    for filename, path in st.session_state.get('generated_files', []):
        with open(path, "rb") as f:
            st.download_button(f"⬇️ Скачать {filename}", f.read(), file_name=filename)

    if st.button("📦 Создать ещё один пакет документов"):
        st.session_state['page'] = 'main_menu'
        st.session_state['form_data'] = {}
        st.session_state['generated_files'] = []